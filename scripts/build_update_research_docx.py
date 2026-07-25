#!/usr/bin/env python3
"""Build the self-contained PatchTicker research report as a polished DOCX."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "UPDATE-TRACKER-RESEARCH-AND-ARCHITECTURE.md"
TS_SOURCE = ROOT / "shared" / "update-feed-item.ts"
SCHEMA_SOURCE = ROOT / "shared" / "update-feed-item.schema.json"
OUTPUT = ROOT / "outputs" / "PatchTicker_Update_Tracker_Research_and_Architecture.docx"

FONT = "Aptos"
MONO = "Aptos Mono"
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20252B"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C8D1DC"
WHITE = "FFFFFF"
GOLD = "B78628"
PORTRAIT_WIDTH_DXA = 9360
LANDSCAPE_WIDTH_DXA = 12960
TABLE_INDENT_DXA = 120


def set_run_font(run, *, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, total_width):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), BORDER)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def add_hyperlink(paragraph, text, url, *, bold=False, italic=False, code=False):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), MONO if code else FONT)
    fonts.set(qn("w:hAnsi"), MONO if code else FONT)
    r_pr.append(fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_inline(paragraph, text, *, base_size=10.5, base_color=INK, shade_code=True):
    for token in filter(None, INLINE_RE.split(text)):
        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2))
            continue
        bold = token.startswith("**") and token.endswith("**")
        code = token.startswith("`") and token.endswith("`")
        italic = token.startswith("*") and token.endswith("*") and not bold
        clean = token[2:-2] if bold else token[1:-1] if code or italic else token
        run = paragraph.add_run(clean)
        set_run_font(run, name=MONO if code else FONT, size=9.5 if code else base_size, color=base_color, bold=bold, italic=italic)
        if code and shade_code:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_GRAY)
            run._element.get_or_add_rPr().append(shading)


def add_numbering_definition(document, num_fmt, text, left=540, hanging=270):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    level.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def style_document(document):
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(document):
    section = document.sections[0]
    configure_section(section)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PATCHTICKER  /  RESEARCH BRIEF")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(64)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("MARKET RESEARCH  •  TELEMETRY  •  ARCHITECTURE")
    set_run_font(run, size=9.5, color=GOLD, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Update Tracker Live Feed")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Research, Telemetry Analytics & Architecture Setup")
    set_run_font(run, size=15, color=DARK_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(48)
    run = meta.add_run("Research date: July 23, 2026  |  Global scope")
    set_run_font(run, size=10.5, color=MUTED)

    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120], PORTRAIT_WIDTH_DXA)
    set_table_borders(table)
    metrics = (("1.08M", "base searches/day"), ("<5 min", "structured-source SLO"), ("99", "5-minute TES"))
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=18, color=NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.5, color=MUTED)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(40)
    note.paragraph_format.space_after = Pt(0)
    r = note.add_run("Prepared for PatchTicker platform implementation")
    set_run_font(r, size=10, color=MUTED, italic=True)
    document.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def column_widths(rows, total_width):
    col_count = max(len(r) for r in rows)
    weights = []
    for idx in range(col_count):
        lengths = [min(len(row[idx]) if idx < len(row) else 0, 55) for row in rows]
        weights.append(max(8, sum(lengths) / max(len(lengths), 1)))
    minimum = 720 if col_count >= 6 else 900
    available = total_width - minimum * col_count
    extra_weight = sum(weights)
    widths = [minimum + round(available * w / extra_weight) for w in weights]
    widths[-1] += total_width - sum(widths)
    return widths


def add_table(document, rows, total_width):
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    widths = column_widths(rows, total_width)
    set_table_geometry(table, widths, total_width)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    font_size = 7.5 if col_count >= 7 else 8 if col_count >= 5 else 8.7
    for r_idx, source_row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            text = source_row[c_idx] if c_idx < len(source_row) else ""
            add_inline(
                p,
                text,
                base_size=font_size,
                base_color=WHITE if r_idx == 0 else INK,
                shade_code=r_idx != 0,
            )
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
            if c_idx > 0 and len(text) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(document, code, language=None):
    keep_together = code.startswith("Provider registry + scheduler")
    if keep_together:
        document.add_page_break()
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [PORTRAIT_WIDTH_DXA], PORTRAIT_WIDTH_DXA)
    cell = table.cell(0, 0)
    if keep_together:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.splitlines() or [""]):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name=MONO, size=7.8, color=INK)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def render_markdown(document, markdown_text, bullet_num_id, decimal_num_id):
    lines = markdown_text.splitlines()
    i = 0
    paragraph_buffer = []
    in_code = False
    code_language = None
    code_lines = []
    last_list_kind = None

    def flush_paragraph():
        nonlocal last_list_kind
        if not paragraph_buffer:
            return
        p = document.add_paragraph()
        p.paragraph_format.keep_together = False
        add_inline(p, " ".join(s.strip() for s in paragraph_buffer))
        paragraph_buffer.clear()
        last_list_kind = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if not in_code:
                in_code = True
                code_language = stripped[3:].strip() or None
                code_lines = []
            else:
                add_code_block(document, "\n".join(code_lines), code_language)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            rows, next_i = parse_table(lines, i)
            wide = len(rows[0]) >= 6
            if wide:
                heading_text = None
                heading_style = None
                if document.paragraphs and document.paragraphs[-1].style.name.startswith("Heading"):
                    heading = document.paragraphs[-1]
                    heading_text = heading.text
                    heading_style = heading.style.name
                    heading._element.getparent().remove(heading._element)
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                if heading_text:
                    moved_heading = document.add_paragraph(style=heading_style)
                    level = int(heading_style.rsplit(" ", 1)[-1])
                    add_inline(moved_heading, heading_text, base_size={1: 16, 2: 13, 3: 11.5}[level], base_color=BLUE if level < 3 else DARK_BLUE)
                add_table(document, rows, LANDSCAPE_WIDTH_DXA)
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=False)
            else:
                add_table(document, rows, PORTRAIT_WIDTH_DXA)
            i = next_i
            last_list_kind = None
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and text.startswith("PatchTicker update-tracker"):
                i += 1
                continue
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline(p, text, base_size={1: 16, 2: 13, 3: 11.5}[level], base_color=BLUE if level < 3 else DARK_BLUE)
            i += 1
            last_list_kind = None
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            if numbered and last_list_kind != "numbered":
                decimal_num_id = add_numbering_definition(document, "decimal", "%1.", 540, 270)
            apply_num(p, bullet_num_id if bullet else decimal_num_id)
            add_inline(p, (bullet or numbered).group(1))
            last_list_kind = "bullet" if bullet else "numbered"
            i += 1
            continue
        if not stripped:
            flush_paragraph()
            last_list_kind = None
            i += 1
            continue
        paragraph_buffer.append(stripped)
        i += 1
    flush_paragraph()


def add_appendix(document, title, source_text, language):
    document.add_page_break()
    p = document.add_paragraph(style="Heading 1")
    add_inline(p, title, base_size=16, base_color=BLUE)
    intro = document.add_paragraph()
    add_inline(intro, "Canonical production artifact included for a self-contained handoff.")
    chunks = source_text.splitlines()
    for start in range(0, len(chunks), 48):
        add_code_block(document, "\n".join(chunks[start:start + 48]), language)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    style_document(document)
    add_cover(document)
    bullet_num_id = add_numbering_definition(document, "bullet", "•", 540, 270)
    decimal_num_id = add_numbering_definition(document, "decimal", "%1.", 540, 270)

    markdown_text = SOURCE.read_text(encoding="utf-8")
    render_markdown(document, markdown_text, bullet_num_id, decimal_num_id)
    add_appendix(document, "Appendix A — TypeScript UpdateFeedItem interface", TS_SOURCE.read_text(encoding="utf-8"), "typescript")
    schema = json.dumps(json.loads(SCHEMA_SOURCE.read_text(encoding="utf-8")), indent=2)
    add_appendix(document, "Appendix B — UpdateFeedItem JSON Schema", schema, "json")

    core = document.core_properties
    core.title = "PatchTicker Update Tracker Live Feed — Research, Telemetry & Architecture"
    core.subject = "Market research, adoption modeling, TES, provider ingestion, normalization, and backend architecture"
    core.author = "PatchTicker"
    core.keywords = "PatchTicker, updates, patch notes, telemetry, SSE, JSON Schema"
    core.comments = "Prepared July 23, 2026"

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
